import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import re
import requests
import base64

# 分层接口设计规范内容
interface_design_note = '''
附录：接口设计分层规范

在软件工程的标准文档体系中，接口设计的描述会根据其抽象层级和设计颗粒度分布在《需求规格书》《概要设计》和《详细设计》三份文档中，具体分布逻辑如下：

---

一、接口设计的层级归属
| 文档类型       | 接口设计内容                                                                 | 示例说明                                                                 |
|--------------------|----------------------------------------------------------------------------------|-----------------------------------------------------------------------------|
| 需求规格书     | 业务接口需求：\n• 系统间交互的业务目标\n• 数据流向及业务规则\n• 非功能性需求（如性能、安全） | “订单系统需对接支付网关，支持每秒处理500+交易请求，响应延迟<200ms”          |
| 概要设计       | 架构级接口规范：\n• 系统/模块间接口拓扑\n• 通信协议与数据格式\n• 接口责任边界              | “用户服务与订单服务采用RESTful API交互，JSON格式，认证使用OAuth2.0协议”     |
| 详细设计       | 实现级接口定义：\n• 精确的API端点/消息格式\n• 字段级数据结构\n• 错误码与状态机            | POST /api/orders { "userId": int, "items": [{"sku":string, "qty":int}] } |

---

二、关键差异对比
| 维度         | 需求规格书                     | 概要设计                         | 详细设计                         |
|------------------|--------------------------------|----------------------------------|----------------------------------|
| 描述焦点     | Why（为什么需要接口）      | What（接口做什么）           | How（接口如何实现）          |
| 颗粒度       | 业务场景级                     | 模块/系统级                      | 代码级                           |
| 输出物       | 用例描述/数据流图              | 接口框图/协议选型                | Swagger文档/IDL定义              |
| 变更影响     | 影响业务目标                   | 影响系统架构                     | 影响具体实现                     |

---

三、常见错误规避
1. 需求与设计混淆  ❌ 错误：在需求文档中定义API参数细节  ✅ 修正：需求文档仅声明“支付接口需支持信用卡和数字货币”，具体支付方式枚举在详细设计中展开  
2. 概要设计缺失关键约束  ❌ 错误：概要设计未说明跨系统事务一致性方案  ✅ 修正：明确标注“订单创建与库存扣减采用SAGA事务模式，通过MQ实现最终一致性”  
3. 详细设计脱离架构  ❌ 错误：详细设计私自将REST协议改为gRPC，违背概要设计约定  ✅ 修正：协议变更需回溯修订概要设计，并评估架构影响  

---

四、敏捷开发中的实践调整
在敏捷项目中（如Scrum），可简化为：
1. 用户故事（需求层）：  “作为商家，我需要商品上下架状态实时同步至搜索服务，延迟不超过5秒”  
2. 技术方案（概要层）：  商品服务 --Kafka--> 搜索服务
3. 任务卡（详细层）：  实现Kafka生产者：  Topic: product-status  Schema: {id:long, status:enum('ONLINE','OFFLINE')}

---

五、接口设计检查清单
| 阶段       | 核心检查项                                                                 |
|----------------|-----------------------------------------------------------------------------|
| 需求规格书 | 是否明确交互方及数据流向？\n• 是否定义SLA（吞吐量/延迟/可用性）？             |
| 概要设计   | 协议选型（HTTP/RPC/MQ）是否合理？\n• 是否规划了接口版本管理策略？               |
| 详细设计   | 字段命名是否遵循团队规范？\n• 错误码体系是否完整？\n• 有无兼容性破坏风险？     |

核心原则：需求层聚焦业务契约，概要层确定技术契约，详细层完成实现契约，三者形成逐级细化的设计闭环。
'''

with open('./README.md', 'r', encoding='utf-8') as f:
    readme = f.read()

sections = [
    ("1. 项目背景与目标", re.search(r'# 2\. 项目意义与作用(.+?)# 3\.', readme, re.S).group(1).strip()),
    ("2. 需求规格", ""),
    ("3. 总体架构设计", ""),
    ("4. 详细设计", ""),
    ("5. 典型业务流程与用例", ""),
    ("6. 运维与监控设计", "")
]

# 提取所有 mermaid 代码块
mermaid_blocks = []
for m in re.finditer(r'```mermaid\s+([\s\S]+?)```', readme):
    mermaid_blocks.append(m.group(1).strip())

# 渲染 mermaid 为图片（用kroki在线API）
def render_mermaid_to_png(mermaid_code, out_path):
    url = 'https://kroki.io/mermaid/png'
    try:
        resp = requests.post(url, data=mermaid_code.encode('utf-8'), timeout=15)
        if resp.status_code == 200:
            with open(out_path, 'wb') as f:
                f.write(resp.content)
            print(f'[OK] Mermaid图已渲染并保存: {out_path}')
            return True
        else:
            print(f'[FAIL] Mermaid渲染失败: {out_path}, 状态码: {resp.status_code}')
    except Exception as e:
        print(f'[ERROR] Mermaid渲染异常: {out_path}, {e}')
    return False

# 生成所有图片，记录路径
img_paths = []
for idx, code in enumerate(mermaid_blocks):
    img_path = f'mermaid_{idx+1}.png'
    if render_mermaid_to_png(code, img_path):
        img_paths.append(img_path)
        # 输出图片文件大小
        if os.path.exists(img_path):
            size = os.path.getsize(img_path)
            print(f'[INFO] 图片 {img_path} 大小: {size} 字节')
        else:
            print(f'[WARN] 图片 {img_path} 未生成')
    else:
        print(f'[WARN] Mermaid源码如下：\n{code}\n')

# 2. 需求规格
req = []
req.append("## 2.1 功能需求\n" + re.search(r'# 3\. 主要功能(.+?)# 4\.', readme, re.S).group(1).strip())
req.append("## 2.2 非功能需求\n- 高可用、可扩展、易维护、低延迟、强安全、支持多云/混合云部署\n- 支持结构化日志、监控、告警、自动健康检查\n- 支持插件化、热插拔、动态扩展\n- 支持多语言、国际化、行业场景化\n")
sections[1] = (sections[1][0], '\n'.join(req))

# 3. 总体架构设计
arch = []
arch_text = re.search(r'# 4\. 软件架构与模块说明(.+?)# 5\.', readme, re.S).group(1).strip()
arch.append(arch_text)
arch.append(re.search(r'# 14\. 系统核心时序图(.+?)# 15\.', readme, re.S).group(1).strip())
sections[2] = (sections[2][0], '\n'.join(arch))

# 4. 详细设计
mod = re.search(r'## 4\.2 主要模块说明(.+?)# 5\.', readme, re.S).group(1).strip()
llm = re.search(r'# 8\. LLM池与元数据管理(.+?)# 9\.', readme, re.S).group(1).strip()
prompt = re.search(r'# 10\. 提示词模板中心(.+?)# 11\.', readme, re.S).group(1).strip()
plugin = re.search(r'# 13\. 插件开发者规范(.+?)# 14\.', readme, re.S).group(1).strip()
api_raw = re.search(r'# 6\. API 说明与接口示例(.+?)# 7\.', readme, re.S).group(1)
api_clean = re.sub(r'- \*\*用法示例\*\*:(.+?)(?=\n- |\n####|\n#|\Z)', '', api_raw, flags=re.S)
api_clean = re.sub(r'- \*\*上层用户流式消费示例.+?```.+?```', '', api_clean, flags=re.S)
api_clean = re.sub(r'```python.+?```', '', api_clean, flags=re.S)
api_clean = re.sub(r'```js.+?```', '', api_clean, flags=re.S)
api_clean = re.sub(r'```json.+?```', '', api_clean, flags=re.S)
api_clean = re.sub(r'\n\s*\n', '\n', api_clean)
api_impl = "## 接口实现\n" + api_clean.strip()
detail = '\n'.join([mod, llm, prompt, plugin, api_impl])
sections[3] = (sections[3][0], detail)

# 5. 典型业务流程与用例
flow = re.search(r'# 14\. 系统核心时序图(.+?)# 15\.', readme, re.S).group(1).strip()
flow = re.sub(r'```python.+?```', '', flow, flags=re.S)
flow = re.sub(r'```js.+?```', '', flow, flags=re.S)
flow = re.sub(r'```json.+?```', '', flow, flags=re.S)
flow = re.sub(r'\n\s*\n', '\n', flow)
sections[4] = (sections[4][0], flow)

# 6. 运维与监控设计
ops = re.search(r'## 5\.4 日志系统(.+?)# 6\.', readme, re.S)
if ops:
    ops_clean = re.sub(r'```.+?```', '', ops.group(1), flags=re.S)
    ops_clean = re.sub(r'- \*\*日志文件\*\*.+?\n', '', ops_clean)
    ops_clean = re.sub(r'- \*\*配置文件\*\*.+?\n', '', ops_clean)
    ops_clean = re.sub(r'- \*\*日志级别\*\*.+?\n', '', ops_clean)
    ops_clean = re.sub(r'- \*\*输出格式\*\*.+?\n', '', ops_clean)
    ops_clean = re.sub(r'\n\s*\n', '\n', ops_clean)
    sections[5] = (sections[5][0], ops_clean.strip())
else:
    sections[5] = (sections[5][0], "详见日志系统章节、监控API、健康检查机制等。")

# 生成Word文档
out_path = 'AICoreDirector_项目前期设计方案.docx'
doc = Document()
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
doc.styles['Normal'].font.size = Pt(11)
doc.add_heading('AICoreDirector 项目前期设计方案', 0)
doc.add_paragraph('（本方案基于项目README自动生成，涵盖需求规格、概要设计、详细设计等，适用于编码前的全局设计与沟通）')

img_idx = 0
for i, (title, content) in enumerate(sections):
    doc.add_heading(title, level=1)
    if i in [2, 4]:
        if img_idx < len(img_paths):
            img_path = img_paths[img_idx]
            size = os.path.getsize(img_path) if os.path.exists(img_path) else 0
            doc.add_paragraph(f'【自动渲染图 {img_idx+1}】 路径: {img_path} 大小: {size} 字节')
            if size > 0:
                doc.add_picture(img_path, width=Inches(5.5))
            else:
                doc.add_paragraph('【图片文件为空或损坏】')
            img_idx += 1
        else:
            doc.add_paragraph('【未检测到可用架构/流程图图片】')
    for part in re.split(r'\n(?=## |# )', content):
        if part.strip():
            if part.strip().startswith('## '):
                doc.add_heading(part.strip().replace('## ', ''), level=2)
            elif part.strip().startswith('# '):
                doc.add_heading(part.strip().replace('# ', ''), level=2)
            else:
                doc.add_paragraph(part.strip())
doc.save(out_path)
print(f'设计方案已生成：{out_path}') 